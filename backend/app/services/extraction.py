"""
Extracts raw text from PDF / DOCX / TXT uploads.
Falls back to OCR (pytesseract + pdf2image) for scanned/image-only PDFs
if those optional packages are installed (Bonus: OCR for Scanned Documents).
"""
import pdfplumber
import docx
from pathlib import Path
from langdetect import detect, LangDetectException


def extract_text(file_path: Path, file_type: str) -> str:
    file_type = file_type.lower()
    if file_type == ".pdf":
        return _extract_pdf(file_path)
    elif file_type == ".docx":
        return _extract_docx(file_path)
    elif file_type == ".txt":
        return file_path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _extract_pdf(file_path: Path) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    text = "\n".join(text_parts).strip()

    if len(text) < 30:  # likely a scanned/image PDF -> try OCR
        text = _ocr_pdf(file_path)
    return text


def _ocr_pdf(file_path: Path) -> str:
    """
    OCR fallback for scanned PDFs.
    Requires:
    - pytesseract
    - pdf2image
    - Tesseract OCR
    - Poppler
    """

    try:
        import pytesseract
        from pdf2image import convert_from_path

        pytesseract.pytesseract.tesseract_cmd = (
            r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        )

        images = convert_from_path(str(file_path))

        ocr_text = []

        for img in images:
            text = pytesseract.image_to_string(img)
            ocr_text.append(text)

        return "\n".join(ocr_text).strip()

    except Exception as e:
        print(f"OCR Error: {e}")
        return ""

def _extract_docx(file_path: Path) -> str:
    document = docx.Document(str(file_path))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


def detect_language(text: str) -> str:
    try:
        return detect(text[:2000])
    except LangDetectException:
        return "en"
