"""
Generates downloadable Risk Assessment Reports in PDF and DOCX formats.
"""
import json
from pathlib import Path
from datetime import datetime

from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle

from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor

from app.core.config import settings
from app.models.models import Document as DocModel, Analysis


def _severity_color(sev: str):
    return {
        "Low": colors.HexColor("#2e7d32"),
        "Medium": colors.HexColor("#f9a825"),
        "High": colors.HexColor("#ef6c00"),
        "Critical": colors.HexColor("#c62828"),
    }.get(sev, colors.grey)


def generate_pdf_report(document: DocModel, analysis: Analysis) -> Path:
    out_path = settings.REPORT_DIR / f"report_doc{document.id}_{int(datetime.utcnow().timestamp())}.pdf"
    doc = SimpleDocTemplate(str(out_path), pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], fontSize=20, spaceAfter=12)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
    body = styles["BodyText"]

    risks = json.loads(analysis.risks_json or "[]")
    elements = [
        Paragraph("AI Contract Risk Assessment Report", title_style),
        Paragraph(f"Document: {document.filename}", body),
        Paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", body),
        Spacer(1, 12),

        Paragraph("Executive Summary", h2),
        Paragraph(analysis.executive_summary or "N/A", body),

        Paragraph("Contract Overview", h2),
        Table([
            ["Contract Type", analysis.contract_type or "N/A"],
            ["Parties", ", ".join(json.loads(analysis.parties or "[]")) or "N/A"],
            ["Effective Date", analysis.effective_date or "N/A"],
            ["Expiry Date", analysis.expiry_date or "N/A"],
            ["Risk Score", f"{analysis.risk_score} / 100"],
            ["Compliance Score", f"{analysis.compliance_score or 0} / 100"],
        ], colWidths=[5 * cm, 10 * cm], style=TableStyle([
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#f0f0f0")),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ])),

        Paragraph("AI Risk Assessment", h2),
    ]

    if risks:
        for r in risks:
            elements.append(Paragraph(
                f"<b><font color='{_severity_color(r.get('severity', 'Low')).hexval()}'>"
                f"[{r.get('severity', 'N/A')}]</font> {r.get('type', '')}</b> "
                f"(confidence: {round(float(r.get('confidence', 0)) * 100)}%)", body))
            elements.append(Paragraph(r.get("explanation", ""), body))
            if r.get("clause_excerpt"):
                elements.append(Paragraph(f"<i>Excerpt: \u201c{r['clause_excerpt']}\u201d</i>", body))
            elements.append(Spacer(1, 6))
    else:
        elements.append(Paragraph("No significant risks detected.", body))

    elements.append(Paragraph("Clause Analysis", h2))
    for label, value in [
        ("Payment Terms", analysis.payment_terms),
        ("Renewal Clause", analysis.renewal_clause),
        ("Confidentiality Clause", analysis.confidentiality_clause),
        ("Termination Clause", analysis.termination_clause),
        ("Responsibilities", analysis.responsibilities),
    ]:
        elements.append(Paragraph(f"<b>{label}:</b> {value or 'Not found in document.'}", body))
        elements.append(Spacer(1, 4))

    elements.append(Paragraph("Recommended Actions", h2))
    for action in json.loads(analysis.recommended_actions or "[]"):
        elements.append(Paragraph(f"&bull; {action}", body))

    doc.build(elements)
    return out_path


def generate_docx_report(document: DocModel, analysis: Analysis) -> Path:
    out_path = settings.REPORT_DIR / f"report_doc{document.id}_{int(datetime.utcnow().timestamp())}.docx"
    d = DocxDocument()

    title = d.add_heading("AI Contract Risk Assessment Report", level=0)
    d.add_paragraph(f"Document: {document.filename}")
    d.add_paragraph(f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    d.add_heading("Executive Summary", level=1)
    d.add_paragraph(analysis.executive_summary or "N/A")

    d.add_heading("Contract Overview", level=1)
    table = d.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    rows = [
        ("Contract Type", analysis.contract_type or "N/A"),
        ("Parties", ", ".join(json.loads(analysis.parties or "[]")) or "N/A"),
        ("Effective Date", analysis.effective_date or "N/A"),
        ("Expiry Date", analysis.expiry_date or "N/A"),
        ("Risk Score", f"{analysis.risk_score} / 100"),
        ("Compliance Score", f"{analysis.compliance_score or 0} / 100"),
    ]
    for k, v in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = k, v

    d.add_heading("AI Risk Assessment", level=1)
    risks = json.loads(analysis.risks_json or "[]")
    if risks:
        for r in risks:
            p = d.add_paragraph()
            run = p.add_run(f"[{r.get('severity')}] {r.get('type')} ")
            run.bold = True
            run.font.color.rgb = RGBColor(0xC6, 0x28, 0x28) if r.get("severity") in ("High", "Critical") else RGBColor(0x33, 0x33, 0x33)
            p.add_run(f"(confidence: {round(float(r.get('confidence', 0)) * 100)}%)")
            d.add_paragraph(r.get("explanation", ""))
            if r.get("clause_excerpt"):
                d.add_paragraph(f"Excerpt: \u201c{r['clause_excerpt']}\u201d").italic = True
    else:
        d.add_paragraph("No significant risks detected.")

    d.add_heading("Clause Analysis", level=1)
    for label, value in [
        ("Payment Terms", analysis.payment_terms),
        ("Renewal Clause", analysis.renewal_clause),
        ("Confidentiality Clause", analysis.confidentiality_clause),
        ("Termination Clause", analysis.termination_clause),
        ("Responsibilities", analysis.responsibilities),
    ]:
        d.add_paragraph(f"{label}: {value or 'Not found in document.'}")

    d.add_heading("Recommended Actions", level=1)
    for action in json.loads(analysis.recommended_actions or "[]"):
        d.add_paragraph(action, style="List Bullet")

    d.save(str(out_path))
    return out_path
